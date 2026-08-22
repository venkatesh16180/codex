import pymupdf   # was: import fitz
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import docx           # pip install python-docx

def extract_pdf(path: str) -> str:
    doc = pymupdf.open(path)   # was: fitz.open(path)
    return '\n'.join(page.get_text() for page in doc)

def extract_epub(path: str) -> str:
    book = epub.read_epub(path)
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), 'html.parser')
        parts.append(soup.get_text())
    return '\n'.join(parts)

def extract_docx(path: str) -> str:
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text for cell in row.cells)
            parts.append(row_text)
    return '\n'.join(parts)

def extract_text_file(path: str) -> str:
    return open(path, encoding='utf-8', errors='ignore').read()

EXTRACTORS = {
    '.pdf': extract_pdf, '.epub': extract_epub, '.docx': extract_docx,
    '.txt': extract_text_file, '.md': extract_text_file,
}